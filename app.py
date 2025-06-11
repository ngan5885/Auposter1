import streamlit as st
st.set_page_config(page_title="AutoPoster", layout="wide")

import pandas as pd
import os
from docx import Document  # Thiếu import này trong code bạn gửi

st.title("📌 AutoPoster - Generate Poster Content from a Topic")

topic_to_dataset = {
    "diabetes": "sample_data/diabetes.csv",
    "air pollution": "sample_data/air_pollution.csv"
}

topic = st.text_input("🔍 Enter your topic", placeholder="e.g., Diabetes")

if topic:
    topic_lower = topic.lower().strip()
    if topic_lower in topic_to_dataset:
        data_path = topic_to_dataset[topic_lower]
        if os.path.exists(data_path):
            df = pd.read_csv(data_path)
            st.success("✅ Dataset loaded successfully.")
            st.subheader("📊 Sample Data Preview:")
            st.dataframe(df.head(10))

            st.subheader("📝 Auto-generated Poster Description:")
            num_rows, num_cols = df.shape
            columns = ', '.join(df.columns[:5]) + "..."

            description = f"""
Topic: {topic.title()}

This study focuses on the issue of {topic.lower()}, a serious environmental challenge caused by emissions from factories, vehicles, and other sources.

The dataset includes {num_rows} key elements:
- Sources of {topic.title()}
- Effects on Human Health
- Proposed Solutions for Pollution Reduction

The analysis aims to raise awareness of the impact of {topic.lower()} and to encourage the adoption of renewable energy and public transportation to mitigate these effects.

This poster provides essential insights that can support future research and promote sustainable environmental practices.
"""

            st.code(description, language="markdown")

            # Tạo file Word
            doc = Document()
            doc.add_heading(f"Poster Content: {topic.title()}", level=1)
            doc.add_paragraph(description)

            word_file_path = f"{topic_lower}_poster_description.docx"
            doc.save(word_file_path)

            # Nút download file Word
            with open(word_file_path, "rb") as file:
                st.download_button(
                    label="📥 Download Description as Word (.docx)",
                    data=file,
                    file_name=word_file_path,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

            # Nút download file txt
            st.download_button(
                label="📥 Download Description as .txt",
                data=description,
                file_name=f"{topic_lower}_poster_description.txt",
                mime="text/plain"
            )

        else:
            st.error("❌ Dataset file not found. Please check your folder.")
    else:
        st.warning("⚠️ No dataset available for this topic yet.")
